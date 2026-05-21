# -*- coding: utf-8 -*-
"""晶体材料样品管理系统 - Flask 主应用"""

import os
import sys
import uuid
import base64
import json
import time
import csv
from io import BytesIO, StringIO
from datetime import datetime
from flask import Flask, request, jsonify, render_template, send_from_directory, send_file, session, redirect, url_for
from werkzeug.exceptions import HTTPException
from functools import wraps
from PIL import Image
import docx
from docx.shared import Inches, Pt

# 当前目录加入 sys.path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import config
import models
import todo_integration
import backup as _backup_module

# 从上级目录导入元素摩尔质量表
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..'))
from molmass_data import elenmentsmasstable

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max upload
app.secret_key = config.SECRET_KEY

# 配置日志输出
import logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(name)s] %(levelname)s: %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S',
)

@app.errorhandler(Exception)
def handle_exception(e):
    # pass through HTTP errors
    if isinstance(e, HTTPException):
        # Return JSON instead of HTML for API routes
        if request.path.startswith('/api/'):
            return jsonify(error=e.description), e.code
        return e
    # now you're handling non-HTTP exceptions only
    app.logger.error(f"Unhandled Exception: {e}", exc_info=True)
    if request.path.startswith('/api/'):
        return jsonify(error="服务器内部错误: " + str(e)), 500
    return "服务器内部错误", 500

# ============================================================
# 登录验证
# ============================================================

@app.before_request
def check_login():
    """所有请求前检查登录状态，白名单除外"""
    allowed = ['login_page', 'do_login', 'static', 'auth_microsoft', 'auth_callback']
    if request.endpoint in allowed:
        return
    if not session.get('logged_in'):
        if request.path.startswith('/api/'):
            return jsonify({'error': '未登录'}), 401
        return redirect(url_for('login_page'))


@app.route('/login')
def login_page():
    return render_template('login.html')


# 登录失败记录: {ip: {'count': int, 'locked_until': float}}
_login_attempts = {}
LOGIN_MAX_FAILURES = 3
LOGIN_LOCK_SECONDS = 86400  # 24小时


@app.route('/api/login', methods=['POST'])
def do_login():
    ip = request.remote_addr or 'unknown'
    now = time.time()

    # 检查是否被锁定
    record = _login_attempts.get(ip, {})
    locked_until = record.get('locked_until', 0)
    if now < locked_until:
        remaining = int((locked_until - now) / 3600) + 1
        return jsonify({'error': f'登录已锁定，请 {remaining} 小时后重试'}), 429

    data = request.get_json()
    password = data.get('password', '') if data else ''

    if password == config.LOGIN_PASSWORD:
        _login_attempts.pop(ip, None)  # 登录成功清除记录
        session['logged_in'] = True
        return jsonify({'success': True})

    # 登录失败
    count = record.get('count', 0) + 1
    if count >= LOGIN_MAX_FAILURES:
        _login_attempts[ip] = {'count': count, 'locked_until': now + LOGIN_LOCK_SECONDS}
        return jsonify({'error': f'密码错误 {LOGIN_MAX_FAILURES} 次，IP 已被锁定 24 小时'}), 429

    _login_attempts[ip] = {'count': count, 'locked_until': 0}
    return jsonify({'error': f'密码错误，还剩 {LOGIN_MAX_FAILURES - count} 次机会'}), 403


@app.route('/api/logout', methods=['POST'])
def do_logout():
    session.clear()
    return jsonify({'success': True})


# ============================================================
# Microsoft To Do OAuth2 路由
# ============================================================

@app.route('/auth/microsoft')
def auth_microsoft():
    """开始 Microsoft OAuth2 授权流程"""
    if not todo_integration.is_configured():
        return jsonify({'error': '请先在 config.py 中配置 MS_CLIENT_ID 和 MS_CLIENT_SECRET'}), 400
    auth_url = todo_integration.get_auth_url()
    return redirect(auth_url)


@app.route('/auth/callback')
def auth_callback():
    """Microsoft OAuth2 回调"""
    code = request.args.get('code')
    error = request.args.get('error')

    if error:
        return f'<h3>授权失败</h3><p>{error}: {request.args.get("error_description", "")}</p><a href="/">返回</a>', 400

    if not code:
        return '<h3>授权失败</h3><p>未收到授权码</p><a href="/">返回</a>', 400

    success, err_msg = todo_integration.acquire_token_by_code(code)
    if success:
        return redirect('/')
    else:
        return f'<h3>授权失败</h3><p>{err_msg}</p><a href="/">返回</a>', 400


@app.route('/api/ms-status', methods=['GET'])
def ms_status():
    """返回 Microsoft To Do 连接状态"""
    return jsonify({
        'configured': todo_integration.is_configured(),
        'connected': todo_integration.is_connected(),
    })


@app.route('/api/ms-disconnect', methods=['POST'])
def ms_disconnect():
    """断开 Microsoft To Do 连接"""
    todo_integration.disconnect()
    return jsonify({'success': True})


# ============================================================
# 页面路由
# ============================================================

@app.route('/')
def index():
    return render_template('index.html')


# ============================================================
# 样品 API
# ============================================================

@app.route('/api/samples', methods=['GET'])
def list_samples():
    query = request.args.get('q', '').strip()
    sort_mode = request.args.get('sort', 'date').strip()
    samples = models.get_all_samples(query if query else None, sort_mode=sort_mode)
    
    # 添加禁用缓存的 Header，防止浏览器缓存旧数据
    response = jsonify(samples)
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    return response

@app.route('/api/samples/sync_todo', methods=['POST'])
def sync_todo():
    """手动或前端自动触发的 To Do 同步"""
    try:
        completed = todo_integration.sync_growing_tasks(models)
        return jsonify({'success': True, 'completed_samples': completed})
    except Exception as e:
        app.logger.error(f"Sync todo failed: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/samples/reorder', methods=['POST'])
def reorder_samples():
    """更新样品手动排序顺序"""
    data = request.get_json()
    if not data or 'ordered_ids' not in data:
        return jsonify({'error': '缺少 ordered_ids 参数'}), 400

    ordered_ids = data['ordered_ids']
    if not isinstance(ordered_ids, list):
        return jsonify({'error': 'ordered_ids 必须是数组'}), 400

    models.reorder_samples(ordered_ids)
    return jsonify({'success': True})


@app.route('/api/samples', methods=['POST'])
def create_sample():
    data = request.get_json()
    if not data or not data.get('id'):
        return jsonify({'error': '样品编号不能为空'}), 400

    # 检查是否已存在
    existing = models.get_sample(data['id'])
    if existing:
        return jsonify({'error': f'样品编号 {data["id"]} 已存在'}), 409

    sample = models.create_sample(data)
    if sample is None:
        return jsonify({'error': '创建失败：未找到刚创建的样品数据'}), 500

    # 新建样品时如果有结束时间，也同步到 To Do
    todo_synced = False
    todo_msg = ''
    sintering_end = data.get('sintering_end', '')
    if sintering_end:
        try:
            ok, msg = todo_integration.create_or_update_todo(data['id'], sintering_end, models, data.get('target_product', ''))
            todo_synced = ok
            todo_msg = msg
        except Exception as e:
            app.logger.error(f'同步 To Do 失败: {e}')
            todo_msg = str(e)

    result = dict(sample) if isinstance(sample, dict) else sample
    result['todo_synced'] = todo_synced
    result['todo_msg'] = todo_msg
    return jsonify(result), 201


@app.route('/api/samples/<sample_id>', methods=['GET'])
def get_sample(sample_id):
    sample = models.get_sample(sample_id)
    if not sample:
        return jsonify({'error': '样品不存在'}), 404
    response = jsonify(sample)
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    return response


@app.route('/api/samples/<sample_id>', methods=['PUT'])
def update_sample(sample_id):
    data = request.get_json()
    if not data:
        return jsonify({'error': '无效数据'}), 400

    existing = models.get_sample(sample_id)
    if not existing:
        return jsonify({'error': '样品不存在'}), 404

    # 如果修改了 ID，检查新 ID 是否已存在
    new_id = data.get('id', sample_id)
    if new_id != sample_id:
        check = models.get_sample(new_id)
        if check:
            return jsonify({'error': f'样品编号 {new_id} 已存在'}), 409

    # 记录旧的 sintering_end 用于后续比较
    old_sintering_end = existing.get('sintering_end', '')

    sample = models.update_sample(sample_id, data)
    if sample is None:
        return jsonify({'error': '更新失败：未能获取更新后的样品数据'}), 500

    # 同步 Microsoft To Do（仅当结束时间有变化时）
    todo_synced = False
    todo_msg = ''
    new_sintering_end = data.get('sintering_end', '')
    if new_sintering_end and new_sintering_end != old_sintering_end:
        try:
            ok, msg = todo_integration.create_or_update_todo(new_id, new_sintering_end, models, data.get('target_product', ''))
            todo_synced = ok
            todo_msg = msg
        except Exception as e:
            app.logger.error(f'同步 To Do 失败: {e}')
            todo_msg = str(e)

    result = dict(sample) if isinstance(sample, dict) else sample
    result['todo_synced'] = todo_synced
    result['todo_msg'] = todo_msg
    return jsonify(result)


@app.route('/api/samples/<sample_id>', methods=['DELETE'])
def delete_sample(sample_id):
    existing = models.get_sample(sample_id)
    if not existing:
        return jsonify({'error': '样品不存在'}), 404

    models.delete_sample(sample_id)
    return jsonify({'success': True})


# ============================================================
# 照片 API
# ============================================================

@app.route('/api/samples/<sample_id>/photos', methods=['POST'])
def upload_photo(sample_id):
    existing = models.get_sample(sample_id)
    if not existing:
        return jsonify({'error': '样品不存在'}), 404

    if 'file' not in request.files:
        return jsonify({'error': '未选择文件'}), 400

    folder = models.get_sample_subfolder(sample_id, 'photos')
    files = request.files.getlist('file')
    uploaded = []
    for file in files:
        if file.filename:
            ext = os.path.splitext(file.filename)[1]
            safe_name = f"{uuid.uuid4().hex}{ext}"
            filepath = os.path.join(folder, safe_name)
            file.save(filepath)
            photo_id = models.add_photo(sample_id, file.filename, filepath)
            _create_thumbnail(filepath)
            uploaded.append({'id': photo_id, 'filename': file.filename, 'filepath': filepath})

    return jsonify(uploaded), 201


@app.route('/api/photos/<int:photo_id>', methods=['DELETE'])
def delete_photo(photo_id):
    models.delete_photo(photo_id)
    return jsonify({'success': True})


# ============================================================
# EDX API
# ============================================================

@app.route('/api/samples/<sample_id>/edx', methods=['POST'])
def upload_edx(sample_id):
    existing = models.get_sample(sample_id)
    if not existing:
        return jsonify({'error': '样品不存在'}), 404

    if 'file' not in request.files:
        return jsonify({'error': '未选择文件'}), 400

    folder = models.get_sample_subfolder(sample_id, 'edx')
    files = request.files.getlist('file')
    uploaded = []
    for file in files:
        if file.filename:
            ext = os.path.splitext(file.filename)[1]
            safe_name = f"{uuid.uuid4().hex}{ext}"
            filepath = os.path.join(folder, safe_name)
            file.save(filepath)
            edx_id = models.add_edx_image(sample_id, file.filename, filepath)
            _create_thumbnail(filepath)
            uploaded.append({'id': edx_id, 'filename': file.filename, 'filepath': filepath})

    return jsonify(uploaded), 201


@app.route('/api/edx/<int:edx_id>/recognize', methods=['POST'])
def recognize_edx(edx_id):
    """调用 GPT Vision API 识别 EDX 谱图"""
    from openai import OpenAI

    app.logger.info("=" * 50)
    app.logger.info(f"[EDX 识别] 开始处理 edx_id={edx_id}")

    conn = models.get_db()
    row = conn.execute("SELECT * FROM edx_images WHERE id = ?", (edx_id,)).fetchone()
    conn.close()

    if not row:
        app.logger.error(f"[EDX 识别] edx_id={edx_id} 不存在于数据库")
        return jsonify({'error': 'EDX 图片不存在'}), 404

    filepath = row['filepath']
    app.logger.info(f"[EDX 识别] 文件路径: {filepath}")

    if not os.path.exists(filepath):
        app.logger.error(f"[EDX 识别] 文件不存在: {filepath}")
        return jsonify({'error': '文件不存在'}), 404

    # 读取图片并 base64 编码
    with open(filepath, 'rb') as f:
        raw_bytes = f.read()
        image_data = base64.b64encode(raw_bytes).decode('utf-8')
    app.logger.info(f"[EDX 识别] 图片已读取, 大小: {len(raw_bytes)} bytes, base64 长度: {len(image_data)}")

    ext = os.path.splitext(filepath)[1].lower()
    mime_map = {'.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                '.bmp': 'image/bmp', '.tif': 'image/tiff', '.tiff': 'image/tiff'}
    mime = mime_map.get(ext, 'image/png')

    try:
        app.logger.info(f"[EDX 识别] 调用 AI API: model={config.OPENAI_MODEL}, base_url={config.OPENAI_BASE_URL}")
        t0 = time.time()

        client = OpenAI(api_key=config.OPENAI_API_KEY, base_url=config.OPENAI_BASE_URL)
        response = client.chat.completions.create(
            model=config.OPENAI_MODEL,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "This is a screenshot of EDX (Energy Dispersive X-ray) quantitative analysis from INCA software. "
                                "Please recognize the table data in the screenshot, including:\n"
                                "1. All element names (e.g. Fe, Te, Ta, etc., i.e. the column headers)\n"
                                "2. Result type (atomic percent or weight percent, check the 'Result Type' dropdown at the bottom)\n"
                                "3. The values for each spectrum (Spectrum 1, Spectrum 2...) for each element\n"
                                "4. The average row in the statistics section\n\n"
                                "Return ONLY a JSON object in the following format, without any other text or markdown markup:\n"
                                '{\n'
                                '  "elements": ["Fe", "Te", "Ta"],\n'
                                '  "result_type": "atomic_percent",\n'
                                '  "spectra": [\n'
                                '    {"label": "Spectrum 1", "values": [2.05, 64.26, 33.70]},\n'
                                '    {"label": "Spectrum 2", "values": [2.29, 63.86, 33.85]}\n'
                                '  ],\n'
                                '  "average": {"label": "Average", "values": [1.68, 64.21, 34.11]}\n'
                                '}\n\n'
                                "Notes:\n"
                                "- The elements array contains all element symbols, corresponding to the table columns\n"
                                "- result_type is either \"atomic_percent\" (atomic percent) or \"weight_percent\" (weight percent)\n"
                                "- Each item in the spectra array has values in the same order as elements\n"
                                "- average is the average row from the statistics section\n"
                                "- Preserve the decimal places as shown in the screenshot\n"
                                "- All labels must be in English (e.g. 'Spectrum 1', 'Average'), even if the screenshot is in another language"
                            )
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime};base64,{image_data}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=2000
        )

        elapsed = time.time() - t0
        app.logger.info(f"[EDX 识别] AI API 响应完成, 耗时: {elapsed:.1f}s")

        result_text = response.choices[0].message.content.strip()
        app.logger.info(f"[EDX 识别] AI 原始返回:\n{result_text}")

        # 尝试解析 JSON
        # 去除可能的 markdown 包裹
        if result_text.startswith('```'):
            lines = result_text.split('\n')
            result_text = '\n'.join(lines[1:-1] if lines[-1].strip() == '```' else lines[1:])
            result_text = result_text.strip()
            app.logger.info(f"[EDX 识别] 去除 markdown 包裹后:\n{result_text}")

        recognized_data = json.loads(result_text)

        # 兼容新旧格式的日志输出
        if isinstance(recognized_data, dict) and 'elements' in recognized_data:
            elements = recognized_data.get('elements', [])
            spectra = recognized_data.get('spectra', [])
            average = recognized_data.get('average', {})
            result_type = recognized_data.get('result_type', 'atomic_percent')
            app.logger.info(f"[EDX 识别] ✅ 解析成功 (INCA 表格格式)")
            app.logger.info(f"  元素: {elements}")
            app.logger.info(f"  结果类型: {result_type}")
            for sp in spectra:
                app.logger.info(f"  {sp.get('label', '?')}: {sp.get('values', [])}")
            if average:
                app.logger.info(f"  {average.get('label', 'Average')}: {average.get('values', [])}")
        else:
            # 旧格式兼容
            app.logger.info(f"[EDX 识别] ✅ 解析成功, 识别到 {len(recognized_data)} 个元素:")
            for item in recognized_data:
                app.logger.info(f"  - {item.get('element', '?')}: wt%={item.get('weight_percent', '?')}, at%={item.get('atomic_percent', '?')}")

        models.update_edx_recognized_data(edx_id, recognized_data)
        app.logger.info(f"[EDX 识别] 数据已保存到数据库")
        app.logger.info("=" * 50)
        return jsonify({'recognized_data': recognized_data})

    except json.JSONDecodeError:
        # GPT 返回的不是有效 JSON，原样返回让前端处理
        app.logger.error(f"[EDX 识别] ❌ JSON 解析失败, AI 返回内容:\n{result_text}")
        models.update_edx_recognized_data(edx_id, None)
        return jsonify({'error': '识别结果格式异常', 'raw': result_text}), 422
    except Exception as e:
        app.logger.error(f"[EDX 识别] ❌ API 调用失败: {e}", exc_info=True)
        models.update_edx_recognized_data(edx_id, None)
        return jsonify({'error': f'GPT API 调用失败: {str(e)}'}), 500


@app.route('/api/edx/<int:edx_id>', methods=['DELETE'])
def delete_edx(edx_id):
    models.delete_edx_image(edx_id)
    return jsonify({'success': True})


@app.route('/api/edx/reorder', methods=['POST'])
def reorder_edx():
    """更新 EDX 图片排序顺序"""
    data = request.get_json()
    if not data or 'ordered_ids' not in data:
        return jsonify({'error': 'Missing ordered_ids'}), 400

    ordered_ids = data['ordered_ids']
    if not isinstance(ordered_ids, list):
        return jsonify({'error': 'ordered_ids must be an array'}), 400

    models.reorder_edx_images(ordered_ids)
    return jsonify({'success': True})


# ============================================================
# XRD API
# ============================================================

@app.route('/api/samples/<sample_id>/xrd', methods=['POST'])
def upload_xrd(sample_id):
    existing = models.get_sample(sample_id)
    if not existing:
        return jsonify({'error': '样品不存在'}), 404

    if 'file' not in request.files:
        return jsonify({'error': '未选择文件'}), 400

    folder = models.get_sample_subfolder(sample_id, 'xrd')
    files = request.files.getlist('file')
    uploaded = []
    for file in files:
        if file.filename:
            ext = os.path.splitext(file.filename)[1]
            safe_name = f"{uuid.uuid4().hex}{ext}"
            filepath = os.path.join(folder, safe_name)
            file.save(filepath)
            xrd_id = models.add_xrd_image(sample_id, file.filename, filepath)
            _create_thumbnail(filepath)
            uploaded.append({'id': xrd_id, 'filename': file.filename, 'filepath': filepath})

    return jsonify(uploaded), 201


@app.route('/api/xrd/<int:xrd_id>', methods=['DELETE'])
def delete_xrd(xrd_id):
    models.delete_xrd_image(xrd_id)
    return jsonify({'success': True})


# ============================================================
# 数据文件 API
# ============================================================

@app.route('/api/samples/<sample_id>/datafiles', methods=['POST'])
def upload_data_file(sample_id):
    existing = models.get_sample(sample_id)
    if not existing:
        return jsonify({'error': '样品不存在'}), 404

    if 'file' not in request.files:
        return jsonify({'error': '未选择文件'}), 400

    folder = models.get_sample_subfolder(sample_id, 'data')
    files = request.files.getlist('file')
    uploaded = []
    for file in files:
        if file.filename:
            ext = os.path.splitext(file.filename)[1]
            safe_name = f"{uuid.uuid4().hex}{ext}"
            filepath = os.path.join(folder, safe_name)
            file.save(filepath)
            file_id = models.add_data_file(sample_id, file.filename, filepath)
            uploaded.append({'id': file_id, 'filename': file.filename, 'filepath': filepath})

    return jsonify(uploaded), 201


@app.route('/api/datafiles/<int:file_id>', methods=['DELETE'])
def delete_data_file(file_id):
    models.delete_data_file(file_id)
    return jsonify({'success': True})


# ============================================================
# 其他文件 API
# ============================================================

@app.route('/api/samples/<sample_id>/otherfiles', methods=['POST'])
def upload_other_file(sample_id):
    existing = models.get_sample(sample_id)
    if not existing:
        return jsonify({'error': '样品不存在'}), 404

    if 'file' not in request.files:
        return jsonify({'error': '未选择文件'}), 400

    folder = models.get_sample_subfolder(sample_id, 'others')
    files = request.files.getlist('file')
    uploaded = []
    for file in files:
        if file.filename:
            ext = os.path.splitext(file.filename)[1]
            safe_name = f"{uuid.uuid4().hex}{ext}"
            filepath = os.path.join(folder, safe_name)
            file.save(filepath)
            file_id = models.add_other_file(sample_id, file.filename, filepath)
            uploaded.append({'id': file_id, 'filename': file.filename, 'filepath': filepath})

    return jsonify(uploaded), 201


@app.route('/api/otherfiles/<int:file_id>', methods=['DELETE'])
def delete_other_file(file_id):
    models.delete_other_file(file_id)
    return jsonify({'success': True})


# ============================================================
# 元素质量计算 API
# ============================================================

@app.route('/api/calculate_mass', methods=['POST'])
def calculate_mass():
    """
    根据元素比例和某一元素的质量，计算其他元素的实际称量质量
    请求体:
    {
        "elements": [{"element": "Fe", "ratio": 1}, {"element": "La", "ratio": 2}, ...],
        "reference_element": "Fe",
        "reference_mass": 0.5
    }
    """
    data = request.get_json()
    if not data:
        return jsonify({'error': '无效数据'}), 400

    elements = data.get('elements', [])
    ref_element = data.get('reference_element', '')
    ref_mass = data.get('reference_mass', 0)

    if not elements or not ref_element or not ref_mass:
        return jsonify({'error': '请填写完整的元素比例和参考元素质量'}), 400

    # 验证参考元素在列表中
    ref_item = None
    for item in elements:
        if item['element'] == ref_element:
            ref_item = item
            break

    if ref_item is None:
        return jsonify({'error': f'参考元素 {ref_element} 不在元素列表中'}), 400

    # 获取摩尔质量
    ref_molar_mass = elenmentsmasstable.get(ref_element)
    if ref_molar_mass is None:
        return jsonify({'error': f'未知元素: {ref_element}'}), 400

    ref_ratio = ref_item['ratio']
    results = []

    for item in elements:
        el = item['element']
        ratio = item['ratio']
        molar_mass = elenmentsmasstable.get(el)

        if molar_mass is None:
            return jsonify({'error': f'未知元素: {el}'}), 400

        if el == ref_element:
            mass = ref_mass
        else:
            # m_B = m_A × (r_B / r_A) × (M_B / M_A)
            mass = ref_mass * (ratio / ref_ratio) * (molar_mass / ref_molar_mass)

        results.append({
            'element': el,
            'ratio': ratio,
            'molar_mass': round(molar_mass, 4),
            'mass': round(mass, 4)
        })

    return jsonify({'results': results})


@app.route('/api/elements', methods=['GET'])
def get_elements():
    """返回所有可用元素列表（带摩尔质量）"""
    return jsonify(elenmentsmasstable)


# ============================================================
# 导出功能
# ============================================================

def _format_element_ratios(element_ratios):
    """将元素比例列表格式化为化学式字符串"""
    if not element_ratios:
        return ''

    # 按元素符号排序，便于阅读
    sorted_ratios = sorted(element_ratios, key=lambda x: x.get('element', ''))
    parts = []
    for item in sorted_ratios:
        element = item.get('element', '')
        ratio = item.get('ratio', 1)
        if ratio == 1:
            parts.append(element)
        else:
            parts.append(f"{element}{ratio}")
    return ''.join(parts)


@app.route('/api/samples/export', methods=['GET'])
def export_samples():
    """导出所有样品为 CSV 格式"""
    lang = request.args.get('lang', 'zh')
    samples = models.get_all_samples()

    output = StringIO()
    writer = csv.writer(output)

    # 写入表头
    if lang == 'en':
        writer.writerow(['Sample ID', 'Element Ratio (Formula)', 'Growth Process', 'Results'])
    else:
        writer.writerow(['编号', '元素比例（化学式）', '生长制度', '结果'])

    # 写入数据
    for sample in samples:
        element_formula = _format_element_ratios(sample.get('element_ratios', []))
        writer.writerow([
            sample.get('id', ''),
            element_formula,
            sample.get('growth_process', ''),
            sample.get('results', '')
        ])

    # 生成响应
    output.seek(0)
    response = send_file(
        BytesIO(output.getvalue().encode('utf-8-sig')),
        mimetype='text/csv; charset=utf-8',
        as_attachment=True,
        download_name=f'samples_export_{config.get_local_now().strftime("%Y%m%d_%H%M%S")}.csv'
    )
    return response


@app.route('/api/samples/<sample_id>/export_word', methods=['GET'])
def export_sample_word(sample_id):
    """导出单个样品为 Word 文档 (中文或英文)"""
    lang = request.args.get('lang', 'zh')
    sample = models.get_sample(sample_id)
    if not sample:
        return jsonify({'error': '样品不存在' if lang == 'zh' else 'Sample not found'}), 404

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        return jsonify({'error': '未安装 python-docx 库，请联系管理员' if lang == 'zh' else 'python-docx library not installed, contact admin'}), 500

    doc = Document()
    
    # 标题
    title_text = f'样品报告: {sample["id"]}' if lang == 'zh' else f'Sample Report: {sample["id"]}'
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 基本信息
    h1_text = '1. 基本信息' if lang == 'zh' else '1. Basic Information'
    doc.add_heading(h1_text, level=1)
    
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    def add_row(key, val):
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(val) if val else ''
        
    add_row('样品编号' if lang == 'zh' else 'Sample ID', sample.get('id', ''))
    add_row('目标产物' if lang == 'zh' else 'Target Product', sample.get('target_product', ''))
    
    status_map_zh = {0: '失败', 1: '成功'}
    status_map_en = {0: 'Fail', 1: 'Success'}
    s_val = sample.get('status', sample.get('is_successful', 2))
    if s_val not in [0, 1]:
        s_val_str = '待定' if lang == 'zh' else 'Pending'
    else:
        s_val_str = status_map_zh.get(s_val, '待定') if lang == 'zh' else status_map_en.get(s_val, 'Pending')
    add_row('状态' if lang == 'zh' else 'Status', s_val_str)
    
    measurements = []
    if sample.get('has_electric'): measurements.append('电学' if lang == 'zh' else 'Electric')
    if sample.get('has_magnetic'): measurements.append('磁性' if lang == 'zh' else 'Magnetic')
    if sample.get('has_xrd'): measurements.append('XRD')
    if sample.get('has_edx'): measurements.append('EDX')
    none_str = '无' if lang == 'zh' else 'None'
    add_row('测量' if lang == 'zh' else 'Measurements', ', '.join(measurements) if measurements else none_str)
    
    # 称重表格
    ratios = sample.get('element_ratios', [])
    masses = sample.get('actual_masses', [])
    if ratios:
        h2_text = '2. 元素比例 & 质量计算' if lang == 'zh' else '2. Element Ratios & Mass'
        doc.add_heading(h2_text, level=1)
        calc_table = doc.add_table(rows=1, cols=4)
        calc_table.style = 'Table Grid'
        hdr = calc_table.rows[0].cells
        hdr[0].text = '元素符号' if lang == 'zh' else 'Symbol'
        hdr[1].text = '摩尔比' if lang == 'zh' else 'Ratio'
        hdr[2].text = '摩尔质量(g/mol)' if lang == 'zh' else 'Molar Mass (g/mol)'
        hdr[3].text = '实际质量(g)' if lang == 'zh' else 'Mass (g)'
        
        for i, item in enumerate(ratios):
            row_cells = calc_table.add_row().cells
            row_cells[0].text = str(item.get('element', ''))
            row_cells[1].text = str(item.get('ratio', ''))
            molar_mass = elenmentsmasstable.get(item.get('element', ''), '')
            row_cells[2].text = str(molar_mass)
            mass_val = ''
            if i < len(masses):
                mass_val = str(masses[i].get('mass', ''))
            row_cells[3].text = mass_val

    # 生长流程
    h3_text = '3. 生长流程' if lang == 'zh' else '3. Growth Process'
    doc.add_heading(h3_text, level=1)
    if sample.get('sintering_start') or sample.get('sintering_end'):
        p = doc.add_paragraph()
        time_label = '时间: ' if lang == 'zh' else 'Time: '
        p.add_run(time_label).bold = True
        to_label = ' 到 ' if lang == 'zh' else ' to '
        p.add_run(f"{sample.get('sintering_start', '—')}{to_label}{sample.get('sintering_end', '—')} ")
        if sample.get('sintering_duration'):
            p.add_run(f"({sample.get('sintering_duration')} h)")
    if sample.get('growth_process'):
        doc.add_paragraph(sample.get('growth_process', ''))
        
    # 结果和备注
    h4_text = '4. 结果与备注' if lang == 'zh' else '4. Results & Notes'
    doc.add_heading(h4_text, level=1)
    if sample.get('results'):
        p = doc.add_paragraph()
        results_label = '结果: ' if lang == 'zh' else 'Results: '
        p.add_run(results_label).bold = True
        doc.add_paragraph(sample.get('results', ''))
    if sample.get('notes'):
        p = doc.add_paragraph()
        notes_label = '备注: ' if lang == 'zh' else 'Notes: '
        p.add_run(notes_label).bold = True
        doc.add_paragraph(sample.get('notes', ''))
        
    # 图片和图表辅助函数
    def add_image_section(section_title, images):
        if not images: return
        doc.add_heading(section_title, level=1)
        for img in images:
            filepath = img.get('filepath')
            if filepath and os.path.exists(filepath):
                ext = os.path.splitext(filepath)[1].lower()
                if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tif', '.tiff']:
                    try:
                        doc.add_picture(filepath, width=Inches(5.0))
                        p = doc.add_paragraph(img.get('filename', ''))
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    except Exception as e:
                        err_msg = '图片加载失败' if lang == 'zh' else 'Image load failed'
                        doc.add_paragraph(f"[{err_msg}: {str(e)}]")
                        
    # 照片
    photos_title = '5. 样品照片' if lang == 'zh' else '5. Photos'
    add_image_section(photos_title, sample.get('photos', []))
    
    # XRD
    xrd_title = '6. XRD 分析' if lang == 'zh' else '6. XRD Analysis'
    add_image_section(xrd_title, sample.get('xrd_images', []))
    
    # EDX
    edx_images = sample.get('edx_images', [])
    if edx_images:
        edx_title = '7. EDX 分析' if lang == 'zh' else '7. EDX Analysis'
        doc.add_heading(edx_title, level=1)
        for img in edx_images:
            r_data = img.get('recognized_data')
            has_table = r_data and isinstance(r_data, dict) and 'elements' in r_data
            
            if not has_table:
                filepath = img.get('filepath')
                if filepath and os.path.exists(filepath):
                    ext = os.path.splitext(filepath)[1].lower()
                    if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tif', '.tiff']:
                        try:
                            doc.add_picture(filepath, width=Inches(5.0))
                        except Exception as e:
                            err_msg = '图片加载失败' if lang == 'zh' else 'Image load failed'
                            doc.add_paragraph(f"[{err_msg}: {str(e)}]")
            
            # 添加 EDX 表格数据
            if has_table:
                elements = r_data.get('elements', [])
                spectra = r_data.get('spectra', [])
                average = r_data.get('average', {})
                result_type = r_data.get('result_type', 'atomic_percent')
                
                res_type_label = '结果类型: ' if lang == 'zh' else 'Result Type: '
                doc.add_paragraph(f"{res_type_label}{result_type}")
                
                if elements and spectra:
                    edx_table = doc.add_table(rows=1, cols=len(elements) + 1)
                    edx_table.style = 'Table Grid'
                    hdr_cells = edx_table.rows[0].cells
                    hdr_cells[0].text = '谱图' if lang == 'zh' else 'Spectrum'
                    for i, el in enumerate(elements):
                        hdr_cells[i+1].text = str(el)
                        
                    for sp in spectra:
                        row_cells = edx_table.add_row().cells
                        label = str(sp.get('label', ''))
                        if lang == 'zh' and label.lower().startswith('spectrum '):
                            label = label.replace('Spectrum ', '谱图 ')
                        row_cells[0].text = label
                        vals = sp.get('values', [])
                        for i in range(len(elements)):
                            row_cells[i+1].text = str(vals[i]) if i < len(vals) else ''
                            
                    if average:
                        row_cells = edx_table.add_row().cells
                        avg_label = '平均' if lang == 'zh' else 'Average'
                        row_cells[0].text = str(average.get('label', avg_label))
                        if lang == 'zh' and row_cells[0].text.lower() == 'average':
                            row_cells[0].text = '平均'
                        vals = average.get('values', [])
                        for i in range(len(elements)):
                            row_cells[i+1].text = str(vals[i]) if i < len(vals) else ''
                            
                doc.add_paragraph() # 空行

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    # 构造导出文件名：编号-目标产物-状态
    def sanitize_filename(name):
        # 允许字母、数字、中文、空格以及 -_.
        return "".join(c if (c.isalnum() or c in '-_. ') else '_' for c in name)

    safe_id = sanitize_filename(sample.get('id', '未命名'))
    
    target_prod = sample.get('target_product', '')
    if not target_prod:
        target_prod = '未知' if lang == 'zh' else 'Unknown'
    safe_target = sanitize_filename(target_prod)
    
    safe_status = sanitize_filename(s_val_str)
    
    download_filename = f"{safe_id}-{safe_target}-{safe_status}.docx"
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=download_filename
    )


# ============================================================
# 静态文件服务 (上传的文件)
# ============================================================

def _create_thumbnail(filepath, max_size=(300, 300)):
    """生成缩略图，保存在同目录，前缀为 thumb_"""
    try:
        # Check if it's a valid image extension
        ext = os.path.splitext(filepath)[1].lower()
        if ext not in ['.jpg', '.jpeg', '.png', '.bmp', '.tif', '.tiff']:
            return

        img = Image.open(filepath)
        # Convert to RGB mode if necessary (e.g. RGBA for PNGs before saving as JPEG, or just to be safe)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
            
        img.thumbnail(max_size)
        
        dir_name = os.path.dirname(filepath)
        base_name = os.path.basename(filepath)
        thumb_path = os.path.join(dir_name, f"thumb_{base_name}")
        
        img.save(thumb_path)
    except Exception as e:
        app.logger.error(f"Failed to create thumbnail for {filepath}: {e}")

@app.route('/uploads/<path:filename>')
def serve_upload(filename):
    use_thumb = request.args.get('thumb') == '1'
    original_filename = None

    # 获取原始文件名
    try:
        conn = models.get_db()
        # 将 URL 路径转为本地操作系统路径格式进行匹配，避免部分反斜杠问题，同时也支持 /
        query_path = f"%{filename.replace('/', os.sep)}"
        for table in ['photos', 'edx_images', 'xrd_images', 'data_files', 'other_files']:
            query = f"SELECT filename FROM {table} WHERE filepath LIKE ?"
            row = conn.execute(query, (query_path,)).fetchone()
            if row:
                original_filename = row['filename']
                break
        conn.close()
    except Exception as e:
        app.logger.error(f"Failed to query original filename: {e}")

    if use_thumb:
        parts = filename.split('/')
        if len(parts) > 0:
            thumb_filename = '/'.join(parts[:-1] + [f"thumb_{parts[-1]}"]) if len(parts) > 1 else f"thumb_{parts[0]}"
            
            # Check if thumb exists (need native path for OS check)
            full_thumb_path = os.path.join(config.UPLOAD_FOLDER, thumb_filename.replace('/', os.sep))
            if os.path.exists(full_thumb_path):
                # Send using POSIX path for Flask security
                if original_filename:
                    thumb_download_name = f"thumb_{original_filename}"
                    return send_from_directory(config.UPLOAD_FOLDER, thumb_filename, download_name=thumb_download_name)
                return send_from_directory(config.UPLOAD_FOLDER, thumb_filename)
                
    if original_filename:
        # 如果是数据文件或普通文件，也可以通过 as_attachment 让用户直接点击链接直接下载，不过前端已经加了 download 属性
        return send_from_directory(config.UPLOAD_FOLDER, filename, download_name=original_filename)

    return send_from_directory(config.UPLOAD_FOLDER, filename)


# ============================================================
# 启动
# ============================================================

if __name__ == '__main__':
    models.init_db()
    # 启动定时备份调度器
    _backup_module.start_scheduler()

    port = getattr(config, 'APP_PORT', 5000)

    # Windows 默认 IPV6_V6ONLY=1，需要 patch 才能让 '::' 同时监听 IPv4
    import socket as _socket
    _orig_bind = _socket.socket.bind

    def _dual_stack_bind(self, address):
        if self.family == _socket.AF_INET6:
            try:
                self.setsockopt(_socket.IPPROTO_IPV6, _socket.IPV6_V6ONLY, 0)
            except (AttributeError, OSError):
                pass
        return _orig_bind(self, address)

    _socket.socket.bind = _dual_stack_bind

    print("=" * 50)
    print("晶体材料样品管理系统")
    print(f"IPv4 访问: http://127.0.0.1:{port}")
    print(f"IPv6 访问: http://[::1]:{port}")
    print("=" * 50)
    app.run(debug=True, host='::', port=port,
            exclude_patterns=['backups/*', 'full_backups/*', 'uploads/*'])
