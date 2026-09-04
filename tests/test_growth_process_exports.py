from csv import reader
from io import BytesIO, StringIO

from docx import Document

import models


def _word_text(response):
    document = Document(BytesIO(response.data))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _csv_rows(response):
    return list(reader(StringIO(response.data.decode("utf-8-sig"))))


def test_word_zh_export_uses_chinese_growth_process_only(client):
    models.create_sample(
        {
            "id": "growth-word-zh",
            "growth_process": "Chinese growth process",
            "growth_process_en": "English growth process",
        }
    )

    response = client.get("/api/samples/growth-word-zh/export_word?lang=zh")

    assert response.status_code == 200
    text = _word_text(response)
    assert "Chinese growth process" in text
    assert "English growth process" not in text


def test_word_en_export_uses_english_growth_process_only(client):
    models.create_sample(
        {
            "id": "growth-word-en",
            "growth_process": "Chinese growth process",
            "growth_process_en": "English growth process",
        }
    )

    response = client.get("/api/samples/growth-word-en/export_word?lang=en")

    assert response.status_code == 200
    text = _word_text(response)
    assert "English growth process" in text
    assert "Chinese growth process" not in text


def test_word_growth_export_does_not_fallback_across_languages(client):
    models.create_sample(
        {
            "id": "growth-word-empty",
            "growth_process": "Chinese-only growth process",
            "growth_process_en": "",
        }
    )

    response = client.get("/api/samples/growth-word-empty/export_word?lang=en")

    assert response.status_code == 200
    assert "Chinese-only growth process" not in _word_text(response)


def test_csv_zh_export_uses_chinese_growth_process_only(client):
    models.create_sample(
        {
            "id": "growth-csv-zh",
            "growth_process": "Chinese CSV process",
            "growth_process_en": "English CSV process",
        }
    )

    response = client.get("/api/samples/export?lang=zh")

    assert response.status_code == 200
    rows = _csv_rows(response)
    row = next(row for row in rows[1:] if row[0] == "growth-csv-zh")
    assert row[2] == "Chinese CSV process"
    assert "English CSV process" not in row


def test_csv_en_export_uses_english_growth_process_only(client):
    models.create_sample(
        {
            "id": "growth-csv-en",
            "growth_process": "Chinese CSV process",
            "growth_process_en": "English CSV process",
        }
    )

    response = client.get("/api/samples/export?lang=en")

    assert response.status_code == 200
    rows = _csv_rows(response)
    row = next(row for row in rows[1:] if row[0] == "growth-csv-en")
    assert row[2] == "English CSV process"
    assert "Chinese CSV process" not in row


def test_csv_growth_export_does_not_fallback_across_languages(client):
    models.create_sample(
        {
            "id": "growth-csv-empty",
            "growth_process": "Chinese-only CSV process",
            "growth_process_en": "",
        }
    )

    response = client.get("/api/samples/export?lang=en")

    assert response.status_code == 200
    rows = _csv_rows(response)
    row = next(row for row in rows[1:] if row[0] == "growth-csv-empty")
    assert row[2] == ""
    assert "Chinese-only CSV process" not in row
