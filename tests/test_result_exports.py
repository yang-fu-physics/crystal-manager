from csv import reader
from io import BytesIO, StringIO

from docx import Document

import models


def _word_text(response):
    document = Document(BytesIO(response.data))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_word_zh_export_uses_saved_chinese_result_only(client):
    models.create_sample(
        {
            "id": "word-zh",
            "results": "中文专属结果",
            "results_en": "English-only result",
        }
    )

    response = client.get("/api/samples/word-zh/export_word?lang=zh")

    assert response.status_code == 200
    assert response.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    text = _word_text(response)
    assert "中文专属结果" in text
    assert "English-only result" not in text


def test_word_en_export_uses_saved_english_result_only(client):
    models.create_sample(
        {
            "id": "word-en",
            "results": "中文专属结果",
            "results_en": "English-only result",
        }
    )

    response = client.get("/api/samples/word-en/export_word?lang=en")

    assert response.status_code == 200
    assert response.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    text = _word_text(response)
    assert "English-only result" in text
    assert "中文专属结果" not in text


def test_word_zh_export_does_not_fallback_to_english_when_chinese_result_is_empty(client):
    models.create_sample(
        {
            "id": "word-zh-empty",
            "results": "",
            "results_en": "English-only result",
        }
    )

    response = client.get("/api/samples/word-zh-empty/export_word?lang=zh")

    assert response.status_code == 200
    assert response.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "English-only result" not in _word_text(response)


def test_word_en_export_does_not_fallback_to_chinese_when_english_result_is_empty(client):
    models.create_sample(
        {
            "id": "word-en-empty",
            "results": "中文专属结果",
            "results_en": "",
        }
    )

    response = client.get("/api/samples/word-en-empty/export_word?lang=en")

    assert response.status_code == 200
    assert response.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "中文专属结果" not in _word_text(response)


def _csv_rows(response):
    return list(reader(StringIO(response.data.decode("utf-8-sig"))))


def test_csv_zh_export_uses_saved_chinese_result_only(client):
    models.create_sample(
        {
            "id": "csv-zh",
            "results": "中文专属结果",
            "results_en": "English-only result",
        }
    )

    response = client.get("/api/samples/export?lang=zh")

    assert response.status_code == 200
    assert response.mimetype == "text/csv"
    rows = _csv_rows(response)
    assert rows[1][3] == "中文专属结果"
    assert sum("中文专属结果" in row for row in rows for _ in [0]) == 1
    assert all("English-only result" not in row for row in rows)


def test_csv_en_export_uses_saved_english_result_only(client):
    models.create_sample(
        {
            "id": "csv-en",
            "results": "中文专属结果",
            "results_en": "English-only result",
        }
    )

    response = client.get("/api/samples/export?lang=en")

    assert response.status_code == 200
    assert response.mimetype == "text/csv"
    rows = _csv_rows(response)
    assert rows[1][3] == "English-only result"
    assert sum("English-only result" in row for row in rows for _ in [0]) == 1
    assert all("中文专属结果" not in row for row in rows)


def test_csv_zh_export_does_not_fallback_to_english_when_chinese_result_is_empty(client):
    models.create_sample(
        {
            "id": "csv-zh-empty",
            "results": "",
            "results_en": "English-only result",
        }
    )

    response = client.get("/api/samples/export?lang=zh")

    assert response.status_code == 200
    assert response.mimetype == "text/csv"
    rows = _csv_rows(response)
    assert rows[1][-1] == ""
    assert all("English-only result" not in row for row in rows)


def test_csv_en_export_does_not_fallback_to_chinese_when_english_result_is_empty(client):
    models.create_sample(
        {
            "id": "csv-en-empty",
            "results": "中文专属结果",
            "results_en": "",
        }
    )

    response = client.get("/api/samples/export?lang=en")

    assert response.status_code == 200
    assert response.mimetype == "text/csv"
    rows = _csv_rows(response)
    assert rows[1][-1] == ""
    assert all("中文专属结果" not in row for row in rows)
